"""
Gera o resumo executivo do Tech Challenge em formato Word (.docx).
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = Path(__file__).resolve().parent.parent
FIG_DIR = BASE_DIR / 'figures'
REPORT_DIR = BASE_DIR / 'reports'
REPORT_DIR.mkdir(parents=True, exist_ok=True)


def set_cell_background(cell, color_hex):
    """Define cor de fundo de uma célula da tabela."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_heading_styled(doc, text, level=1, color='1F4E79'):
    """Adiciona heading com cor customizada."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return h


def add_paragraph_justified(doc, text):
    """Adiciona parágrafo justificado."""
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_table_two_col(doc, dados, col1_header='', col2_header='',
                       header_bg='1F4E79', alternate=True):
    """Adiciona tabela de 2 colunas com cabeçalho colorido."""
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'

    # Cabeçalho
    hdr = table.rows[0].cells
    hdr[0].text = col1_header
    hdr[1].text = col2_header
    for cell in hdr:
        set_cell_background(cell, header_bg)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string('FFFFFF')
                run.font.size = Pt(11)

    # Linhas
    for i, (k, v) in enumerate(dados):
        row = table.add_row().cells
        row[0].text = str(k)
        row[1].text = str(v)
        if alternate and i % 2 == 0:
            for cell in row:
                set_cell_background(cell, 'F2F2F2')
        for cell in row:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table


def gerar_documento():
    doc = Document()

    # ---- Estilos base ----
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Margens
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ============================================================
    # CAPA
    # ============================================================
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run('Tech Challenge — Fase 1')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string('1F4E79')

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitulo.add_run('Case NPS Preditivo — Resumo Executivo')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string('595959')

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('\nPós-Tech AI Scientist · Metodologia CRISP-DM\n').italic = True

    doc.add_paragraph()

    # ============================================================
    # 1. CONTEXTO
    # ============================================================
    add_heading_styled(doc, '1. Contexto do Problema', 1)
    add_paragraph_justified(doc,
        'Um e-commerce em rápida expansão observou alta variabilidade no Net '
        'Promoter Score (NPS) entre clientes com perfis operacionais aparentemente '
        'similares. Como o NPS é coletado apenas APÓS a jornada de compra, a '
        'empresa só descobre que o cliente está insatisfeito DEPOIS que o '
        'problema já ocorreu — perdendo a chance de agir preventivamente.'
    )
    add_paragraph_justified(doc,
        'O objetivo deste trabalho é, a partir de dados operacionais (pedidos, '
        'logística e atendimento), entender quais fatores realmente influenciam '
        'a satisfação e propor uma forma de antecipar o NPS antes da pesquisa '
        'formal — abrindo espaço para ações preventivas.'
    )

    # Por que NPS importa
    add_heading_styled(doc, '1.1. Por que o NPS importa em e-commerce', 2)
    p = doc.add_paragraph()
    p.add_run('Recompra: ').bold = True
    p.add_run('promotores recompram em ritmo muito superior ao de detratores, '
              'sendo a métrica mais ligada ao Lifetime Value (LTV).')
    p = doc.add_paragraph()
    p.add_run('Boca a boca: ').bold = True
    p.add_run('em e-commerce, avaliações públicas em marketplaces e redes sociais '
              'são decisivas na conversão de novos compradores.')
    p = doc.add_paragraph()
    p.add_run('Market share: ').bold = True
    p.add_run('detratores ativos podem afastar prospects e corroer participação '
              'de mercado de forma orgânica.')

    # ============================================================
    # 2. DADOS
    # ============================================================
    doc.add_page_break()
    add_heading_styled(doc, '2. Visão Geral dos Dados', 1)
    add_paragraph_justified(doc,
        'A base contém 2.500 registros, cada um representando um cliente único '
        'com seu pedido associado. Foram analisadas 18 variáveis descritivas '
        'mais o NPS Score como variável-alvo.'
    )

    add_heading_styled(doc, '2.1. Qualidade dos dados', 2)
    add_table_two_col(doc, [
        ('Linhas', '2.500'),
        ('Colunas', '19'),
        ('Valores nulos', '0'),
        ('Linhas duplicadas', '0'),
        ('Clientes únicos', '2.500 (1 cliente = 1 pedido)'),
    ], col1_header='Verificação', col2_header='Resultado')

    p = doc.add_paragraph()
    p.add_run('\nConclusão: ').bold = True
    p.add_run('a base é consistente — não foi necessário tratamento de nulos, '
              'duplicados ou inconsistências de tipo. Apenas foi adicionada a '
              'classificação NPS (Detrator / Neutro / Promotor) e features '
              'derivadas para a modelagem.')

    add_heading_styled(doc, '2.2. Achado relevante: NPS é contínuo', 2)
    add_paragraph_justified(doc,
        'Apesar do dicionário descrever o NPS como "variando de 0 a 10", os '
        'valores são contínuos com casas decimais (ex.: 7.3, 8.6). Isso sugere '
        'que a variável não é uma resposta direta de pesquisa tradicional, mas '
        'um score internamente ponderado pela empresa. A classificação em '
        'Detrator/Neutro/Promotor foi feita usando arredondamento para baixo '
        '(floor), de forma conservadora.'
    )

    # ============================================================
    # 3. PRINCIPAIS INSIGHTS
    # ============================================================
    doc.add_page_break()
    add_heading_styled(doc, '3. Principais Insights da EDA', 1)

    # 3.1 Distribuição
    add_heading_styled(doc, '3.1. Distribuição da base', 2)
    add_table_two_col(doc, [
        ('Detratores (NPS 0-6)', '84,4% — situação crítica'),
        ('Neutros (NPS 7-8)', '11,2%'),
        ('Promotores (NPS 9-10)', '4,4%'),
        ('NPS Score oficial', '-80,0 (extremamente negativo)'),
    ], col1_header='Categoria', col2_header='Resultado')

    add_paragraph_justified(doc,
        '\nQualquer NPS acima de 0 já é considerado positivo. Acima de 50 é '
        'tido como excelente. O valor encontrado (-80) representa um cenário '
        'de ruptura severa na experiência do cliente — para cada cliente '
        'promotor há aproximadamente 19 detratores.'
    )

    if (FIG_DIR / '02_classificacao_nps.png').exists():
        doc.add_picture(str(FIG_DIR / '02_classificacao_nps.png'), width=Inches(4.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3.2 Drivers
    add_heading_styled(doc, '3.2. O que mais impacta o NPS', 2)
    add_paragraph_justified(doc,
        'Análise de correlação revelou os principais drivers — variáveis que '
        'mais "puxam" o NPS para cima ou para baixo:'
    )
    add_table_two_col(doc, [
        ('Recompra em 30 dias', '+0,60 (mais forte)'),
        ('CSAT interno', '+0,57'),
        ('Atraso na entrega', '-0,60 (mais forte negativo)'),
        ('Nº de reclamações', '-0,50'),
        ('Contatos com SAC', '-0,35'),
        ('Tempo de resolução', '-0,19'),
    ], col1_header='Variável', col2_header='Correlação com NPS')

    if (FIG_DIR / '05_correlacao_nps.png').exists():
        doc.add_picture(str(FIG_DIR / '05_correlacao_nps.png'), width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3.3 Atraso
    add_heading_styled(doc, '3.3. Ponto de ruptura: atraso na entrega', 2)
    add_paragraph_justified(doc,
        'O atraso na entrega é o pior destruidor de NPS. Com zero dias de '
        'atraso, o NPS médio fica próximo de 6,9 (zona neutra). A cada dia '
        'adicional de atraso, o NPS cai consistentemente.'
    )
    if (FIG_DIR / '03_nps_por_atraso.png').exists():
        doc.add_picture(str(FIG_DIR / '03_nps_por_atraso.png'), width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3.4 Outros
    add_heading_styled(doc, '3.4. Outros achados', 2)
    p = doc.add_paragraph()
    p.add_run('Regiões: ').bold = True
    p.add_run('todas as cinco regiões do Brasil têm NPS médio muito similar '
              '(~4,2 a 4,5). O problema é sistêmico, não regional.')
    p = doc.add_paragraph()
    p.add_run('Demografia: ').bold = True
    p.add_run('idade do cliente e tempo de relacionamento (tenure) têm impacto '
              'baixo no NPS. O problema não é de perfil — é operacional.')
    p = doc.add_paragraph()
    p.add_run('Recompra: ').bold = True
    p.add_run('100% dos promotores recompram em 30 dias; quase nenhum detrator '
              'recompra. Recompra é tanto consequência quanto preditor.')

    # ============================================================
    # 4. MODELO PREDITIVO
    # ============================================================
    doc.add_page_break()
    add_heading_styled(doc, '4. Modelo Preditivo Proposto', 1)
    add_paragraph_justified(doc,
        'Para antecipar o NPS antes da pesquisa, foram construídos dois modelos '
        'complementares:'
    )

    add_heading_styled(doc, '4.1. Modelo de Regressão (estima a nota de NPS)', 2)
    add_paragraph_justified(doc,
        'Prediz a nota de NPS em escala contínua de 0 a 10. Útil para gerar um '
        'score de risco ranqueável — permite priorizar clientes em risco do mais '
        'crítico ao menos crítico.'
    )
    add_table_two_col(doc, [
        ('Modelo', 'Regressão Linear (melhor por R²)'),
        ('R² (variância explicada)', '0,65 — explica 65% da variação do NPS'),
        ('MAE (erro médio absoluto)', '1,19 pontos de NPS'),
        ('Uso recomendado', 'Ranking de clientes por score de risco'),
    ], col1_header='Métrica', col2_header='Valor')

    add_heading_styled(doc, '4.2. Modelo de Classificação (Detrator vs Não)', 2)
    add_paragraph_justified(doc,
        'Prediz se o cliente será detrator (1) ou não (0). Útil para gatilho '
        'automático de ações preventivas (cupom, contato proativo, priorização '
        'de SAC).'
    )
    add_table_two_col(doc, [
        ('Modelo', 'Random Forest (melhor por F1)'),
        ('Accuracy', '0,91'),
        ('Precision', '0,91 — 91% dos alertas são detratores reais'),
        ('Recall', '~1,00 — captura praticamente todos os detratores'),
        ('AUC-ROC', '0,92 — excelente poder discriminativo'),
    ], col1_header='Métrica', col2_header='Valor')

    if (FIG_DIR / '10_feature_importance_classificacao.png').exists():
        doc.add_picture(str(FIG_DIR / '10_feature_importance_classificacao.png'),
                        width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph('Variáveis mais importantes para o modelo de classificação')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)

    add_heading_styled(doc, '4.3. Como o modelo seria usado na prática', 2)
    add_paragraph_justified(doc,
        'A solução pode ser integrada ao sistema de pedidos da seguinte forma:'
    )
    p = doc.add_paragraph(style='List Number')
    p.add_run('Logo após a confirmação do pedido, o modelo gera um score inicial '
              'baseado no perfil do cliente e características do pedido.')
    p = doc.add_paragraph(style='List Number')
    p.add_run('Conforme eventos da jornada ocorrem (preparação, expedição, '
              'entrega, atendimento), o score é atualizado em tempo real.')
    p = doc.add_paragraph(style='List Number')
    p.add_run('Quando o score cruza um limiar de risco, dispara fluxo automático: '
              'priorização logística, contato proativo do SAC ou cupom de '
              'desculpas.')
    p = doc.add_paragraph(style='List Number')
    p.add_run('A pesquisa formal de NPS valida o modelo e alimenta o '
              're-treinamento periódico (mensal recomendado).')

    # ============================================================
    # 5. RECOMENDAÇÕES
    # ============================================================
    doc.add_page_break()
    add_heading_styled(doc, '5. Recomendações para o Negócio', 1)
    add_paragraph_justified(doc,
        'A análise levou a quatro recomendações priorizadas por impacto e '
        'esforço de implementação:'
    )

    recomendacoes = [
        ('1. Logística como prioridade #1',
         'Atraso é o pior driver de NPS. Revisar SLA com transportadoras, '
         'sub-prometer e super-entregar nos prazos, e instituir alerta interno '
         'automático quando um pedido cruza D-1 antes do prazo prometido. '
         'Impacto alto, esforço médio.'),
        ('2. Score de risco em produção',
         'Integrar modelo de classificação à plataforma de pedidos para gerar '
         'score de risco em tempo real. Threshold inicial em 0,5; calibrar '
         'mensalmente. Impacto alto, esforço baixo.'),
        ('3. SAC priorizado por risco',
         'Reordenar fila de tickets pelo score de risco do cliente e por número '
         'de contatos prévios. Cliente que contatou 2+ vezes já é forte sinal '
         'de detração. Impacto médio, esforço médio.'),
        ('4. Programa de recompra agressivo',
         'Recompra em 30 dias é o sinal positivo mais forte. Incentivar '
         'segunda compra com cupom direcionado e cadência de e-mail nos '
         'primeiros 14 dias após a entrega. Impacto médio, esforço baixo.'),
    ]
    for titulo, texto in recomendacoes:
        p = doc.add_paragraph()
        run = p.add_run(titulo)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string('1F4E79')
        run.font.size = Pt(12)
        add_paragraph_justified(doc, texto)

    # ============================================================
    # 6. LIMITAÇÕES
    # ============================================================
    doc.add_page_break()
    add_heading_styled(doc, '6. Limitações e Riscos', 1)
    add_paragraph_justified(doc,
        'É essencial que stakeholders conheçam as limitações desta análise '
        'antes de usar os resultados em decisões de larga escala.'
    )

    limitacoes = [
        ('Tamanho da base',
         '2.500 clientes é uma amostra pequena para modelos complexos. '
         'Recomenda-se ampliar a base antes de testes A/B em produção.'),
        ('Vazamento de rótulo',
         'Variáveis como csat_internal_score e complaints_count podem ser '
         'coletadas APÓS o evento que gera detração — usar em produção exige '
         'validar quais features estão de fato disponíveis ANTES da pesquisa NPS.'),
        ('Variáveis reativas vs preventivas',
         'Reclamações e contatos com SAC já indicam que algo deu errado — usar '
         'esses sinais para "prever" detratores é tardio. Para prevenção real, '
         'priorizar variáveis que precedem a manifestação do cliente, como '
         'delivery_delay_days e delivery_time_days.'),
        ('Sem dimensão temporal',
         'A base não tem datas, então não modelamos tendências, sazonalidade '
         'ou efeitos de eventos pontuais (Black Friday, Natal, greves logísticas).'),
        ('NPS contínuo (não inteiro)',
         'O nps_score é contínuo e não corresponde estritamente a uma '
         'pesquisa NPS tradicional. A interpretação como Detrator/Neutro/Promotor '
         'mantém validade conceitual, mas a empresa deve esclarecer como esse '
         'score foi originalmente construído.'),
    ]
    for titulo, texto in limitacoes:
        p = doc.add_paragraph()
        run = p.add_run('• ' + titulo + ': ')
        run.bold = True
        p.add_run(texto)

    # ============================================================
    # 7. CONCLUSÃO
    # ============================================================
    add_heading_styled(doc, '7. Conclusão', 1)
    add_paragraph_justified(doc,
        'A análise mostrou que a empresa enfrenta um cenário crítico de '
        'satisfação (NPS ≈ -80, com 84% de detratores), mas que esse cenário '
        'é majoritariamente operacional — não regional nem demográfico. O '
        'principal vetor de detração é o atraso na entrega, seguido por '
        'reclamações e contatos com SAC.'
    )
    add_paragraph_justified(doc,
        'Os modelos preditivos construídos têm performance suficiente para uso '
        'prático: o modelo de classificação atinge F1 de 0,95 e AUC de 0,92 — '
        'patamar de excelência — e pode ser integrado ao sistema de pedidos '
        'para gerar alertas em tempo real, transformando uma métrica reativa '
        'em uma alavanca preventiva.'
    )
    add_paragraph_justified(doc,
        'O combo proposto — classificação para acionamento + regressão para '
        'ranqueamento — entrega à empresa um ferramental para mover do reativo '
        'ao proativo na gestão da experiência do cliente, com foco prioritário '
        'em logística como primeira frente de melhoria.'
    )

    # Salvar
    output_path = REPORT_DIR / 'resumo_executivo.docx'
    doc.save(output_path)
    print(f'Documento salvo: {output_path}')
    return output_path


if __name__ == '__main__':
    gerar_documento()
