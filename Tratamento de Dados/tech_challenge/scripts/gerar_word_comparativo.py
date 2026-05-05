"""
Gera o documento Word com o comparativo V1 vs V2.
"""

from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = Path(__file__).resolve().parent.parent
FIG_DIR = BASE_DIR / 'figures_v2'
FIG_V1_DIR = BASE_DIR / 'figures'
REPORT_DIR = BASE_DIR / 'reports'


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color='1F4E79'):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return h


def add_para_just(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_three_col_table(doc, dados, headers, header_bg='1F4E79'):
    """Tabela de 3 colunas (métrica, v1, v2) com cabeçalho colorido."""
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], header_bg)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor.from_string('FFFFFF')
                r.font.size = Pt(11)

    for i, linha in enumerate(dados):
        row = table.add_row().cells
        for j, val in enumerate(linha):
            row[j].text = str(val)
        if i % 2 == 0:
            for cell in row:
                set_cell_bg(cell, 'F2F2F2')
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table


def add_four_col_table(doc, dados, headers, header_bg='1F4E79'):
    """Tabela de 4 colunas."""
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], header_bg)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor.from_string('FFFFFF')
                r.font.size = Pt(11)

    for i, linha in enumerate(dados):
        row = table.add_row().cells
        for j, val in enumerate(linha):
            row[j].text = str(val)
        if i % 2 == 0:
            for cell in row:
                set_cell_bg(cell, 'F2F2F2')
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table


def gerar():
    # Carrega dados do comparativo
    df_geral = pd.read_csv(REPORT_DIR / 'comparativo_visao_geral.csv')
    df_reg = pd.read_csv(REPORT_DIR / 'comparativo_regressao.csv')
    df_cls = pd.read_csv(REPORT_DIR / 'comparativo_classificacao.csv')
    df_corr = pd.read_csv(REPORT_DIR / 'comparativo_correlacoes.csv', index_col=0)

    geral = {row['metrica']: (float(row['v1']), float(row['v2']))
             for _, row in df_geral.iterrows()}

    # Modelos: tem linhas com 'modelo_v1'/'modelo_v2' que são strings — separar
    def ler_metricas_modelo(df):
        out = {}
        for _, row in df.iterrows():
            k = row['metrica']
            v1, v2 = row['v1'], row['v2']
            if k.startswith('modelo'):
                out[k] = (str(v1), str(v2))
            else:
                out[k] = (float(v1), float(v2))
        return out

    reg = ler_metricas_modelo(df_reg)
    cls = ler_metricas_modelo(df_cls)

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ==================== CAPA ====================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Tech Challenge — Fase 1')
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string('1F4E79')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Comparativo V1 vs V2 — Tratamento de NPS=0 como outliers')
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string('595959')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('\nPós-Tech AI Scientist · Metodologia CRISP-DM\n').italic = True
    doc.add_paragraph()

    # ==================== 1. CONTEXTO ====================
    add_heading(doc, '1. Contexto da Análise', 1)
    add_para_just(doc,
        'Na versão 1 (V1) do trabalho, mantivemos toda a base de 2.500 clientes, '
        'incluindo 158 registros com NPS=0. A análise da V1 mostrou que esses '
        'clientes têm perfil de detratores genuínos (mais atrasos, mais '
        'reclamações, mais contatos com SAC), e portanto foram mantidos.'
    )
    add_para_just(doc,
        'Esta versão 2 (V2) refaz toda a análise removendo os NPS=0 como '
        'outliers, com o objetivo de avaliar:'
    )
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Se as ').italic = False
    p.add_run('correlações entre variáveis e NPS').bold = True
    p.add_run(' mudam de forma relevante')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Se a ').italic = False
    p.add_run('performance dos modelos preditivos').bold = True
    p.add_run(' melhora, piora ou se mantém')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Se o ').italic = False
    p.add_run('storytelling para o negócio').bold = True
    p.add_run(' é alterado e qual versão deve ser adotada como definitiva')

    # ==================== 2. METODOLOGIA ====================
    add_heading(doc, '2. Metodologia', 1)
    add_para_just(doc,
        'Para garantir comparabilidade, a V2 utiliza exatamente o mesmo pipeline '
        'da V1: mesma engenharia de features (severidade de problemas, percentual '
        'de desconto), mesmas variáveis de entrada, mesmo split treino/teste '
        '(80/20 com seed fixa), mesmos algoritmos avaliados (Regressão Linear e '
        'Random Forest para regressão; Regressão Logística e Random Forest para '
        'classificação) e mesmas métricas. A única diferença é a remoção dos 158 '
        'registros com NPS=0.'
    )

    # ==================== 3. VISÃO GERAL ====================
    doc.add_page_break()
    add_heading(doc, '3. Visão Geral — Distribuição da Base', 1)

    delta_linhas = int(geral['n_linhas'][1]) - int(geral['n_linhas'][0])
    delta_nps = geral['nps_medio'][1] - geral['nps_medio'][0]
    delta_oficial = geral['nps_oficial'][1] - geral['nps_oficial'][0]
    delta_det = geral['pct_detrator'][1] - geral['pct_detrator'][0]
    delta_neu = geral['pct_neutro'][1] - geral['pct_neutro'][0]
    delta_pro = geral['pct_promotor'][1] - geral['pct_promotor'][0]

    add_four_col_table(doc, [
        ['Tamanho da base', f'{int(geral["n_linhas"][0]):,}',
         f'{int(geral["n_linhas"][1]):,}', f'{delta_linhas:+,}'],
        ['NPS médio', f'{geral["nps_medio"][0]:.2f}',
         f'{geral["nps_medio"][1]:.2f}', f'{delta_nps:+.2f}'],
        ['NPS oficial', f'{geral["nps_oficial"][0]:.1f}',
         f'{geral["nps_oficial"][1]:.1f}', f'{delta_oficial:+.2f}'],
        ['% Detratores', f'{geral["pct_detrator"][0]:.1f}%',
         f'{geral["pct_detrator"][1]:.1f}%', f'{delta_det:+.2f} pp'],
        ['% Neutros', f'{geral["pct_neutro"][0]:.1f}%',
         f'{geral["pct_neutro"][1]:.1f}%', f'{delta_neu:+.2f} pp'],
        ['% Promotores', f'{geral["pct_promotor"][0]:.1f}%',
         f'{geral["pct_promotor"][1]:.1f}%', f'{delta_pro:+.2f} pp'],
    ], headers=['Métrica', 'V1 (com NPS=0)', 'V2 (sem NPS=0)', 'Δ'])

    add_para_just(doc, '')
    p = doc.add_paragraph()
    p.add_run('Leitura: ').bold = True
    p.add_run(f'a remoção de {abs(delta_linhas)} registros (6,3% da base) eleva '
              f'o NPS médio em apenas 0,30 ponto e o NPS oficial em 1,35 ponto. '
              'Mesmo retirando os clientes mais insatisfeitos, o NPS oficial '
              'continua em -78,6 — situação ainda crítica. Isso ')
    p.add_run('confirma que o problema não está concentrado nos "extremos"').italic = True
    p.add_run(', mas espalhado por toda a base.')

    # Figura comparativa de distribuição
    if (FIG_DIR / '04_comparativo_distribuicao.png').exists():
        doc.add_picture(str(FIG_DIR / '04_comparativo_distribuicao.png'),
                        width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ==================== 4. CORRELAÇÕES ====================
    doc.add_page_break()
    add_heading(doc, '4. Mudança nas Correlações', 1)
    add_para_just(doc,
        'Esta é uma das comparações mais reveladoras: como as correlações entre '
        'cada variável e o NPS se alteram quando os detratores extremos são '
        'removidos.'
    )

    # Top 6 mudanças
    top6 = df_corr.head(6)
    linhas_corr = []
    for var, row in top6.iterrows():
        c1, c2, delta = row['v1'], row['v2'], row['delta']
        if (c1 > 0 and c2 < 0) or (c1 < 0 and c2 > 0):
            comentario = '⚠ inverteu sinal'
        elif abs(c2) < abs(c1):
            comentario = 'enfraqueceu'
        else:
            comentario = 'fortaleceu'
        linhas_corr.append([var, f'{c1:+.3f}', f'{c2:+.3f}', f'{delta:+.3f} ({comentario})'])

    add_four_col_table(doc, linhas_corr,
                       headers=['Variável', 'V1', 'V2', 'Δ'])

    add_para_just(doc, '')
    p = doc.add_paragraph()
    p.add_run('Observações principais: ').bold = True

    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Nenhuma correlação inverteu de sinal').bold = True
    p.add_run(' — as relações estruturais entre as variáveis e o NPS são as '
              'mesmas. Atraso continua sendo o pior driver, recompra continua '
              'sendo o melhor.')

    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Quase todas as correlações enfraqueceram').bold = True
    p.add_run(' — o que faz sentido: os clientes com NPS=0 são justamente '
              'aqueles cuja "narrativa" é mais nítida (atraso → reclamação → '
              'NPS zero). Removê-los suaviza os extremos do sinal.')

    p = doc.add_paragraph(style='List Bullet')
    p.add_run('A única correlação que fortaleceu de forma relevante foi '
              'repeat_purchase_30d').bold = True
    p.add_run(' (+0.570 → +0.601). Quem dá nota 0 raramente recompra, mas '
              'alguns dão nota baixa por outros motivos circunstanciais — sem '
              'eles, a relação "recompra ↔ promoção" fica mais limpa.')

    # Figura comparativa de correlações
    if (FIG_DIR / '05_comparativo_correlacoes.png').exists():
        doc.add_picture(str(FIG_DIR / '05_comparativo_correlacoes.png'),
                        width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ==================== 5. MODELOS ====================
    doc.add_page_break()
    add_heading(doc, '5. Performance dos Modelos', 1)

    add_heading(doc, '5.1. Modelo de Regressão (estima nota de NPS)', 2)
    add_four_col_table(doc, [
        ['MAE (erro absoluto)',  f'{reg["MAE"][0]:.3f}',  f'{reg["MAE"][1]:.3f}',  f'{reg["MAE"][1]-reg["MAE"][0]:+.3f}'],
        ['RMSE',                 f'{reg["RMSE"][0]:.3f}', f'{reg["RMSE"][1]:.3f}', f'{reg["RMSE"][1]-reg["RMSE"][0]:+.3f}'],
        ['R² (variância expl.)', f'{reg["R²"][0]:.3f}',   f'{reg["R²"][1]:.3f}',   f'{reg["R²"][1]-reg["R²"][0]:+.3f}'],
        ['Melhor algoritmo',     reg['modelo_v1'][0],     reg['modelo_v2'][1],     '—'],
    ], headers=['Métrica', 'V1', 'V2', 'Δ'])

    add_para_just(doc, '')
    p = doc.add_paragraph()
    p.add_run('Leitura: ').bold = True
    p.add_run('o erro absoluto médio (MAE) e o RMSE caem ligeiramente na V2 — '
              'o modelo erra um pouco menos. Porém, o R² ')
    p.add_run('cai').bold = True
    p.add_run(' de 0,65 para 0,57 (-0,08). Isso parece contraditório, mas tem '
              'explicação clara: ao remover os NPS=0, a variância TOTAL do '
              'target diminui (a base ficou mais homogênea), e o modelo passa '
              'a explicar uma fatia menor dessa variância — mesmo com erros '
              'absolutos menores. Em outras palavras: ')
    p.add_run('a V1 tinha "mais coisa" para explicar, e fazia isso bem; a V2 '
              'tem menos a explicar.').italic = True

    add_heading(doc, '5.2. Modelo de Classificação (Detrator vs Não)', 2)
    add_four_col_table(doc, [
        ['Accuracy',  f'{cls["Accuracy"][0]:.3f}',  f'{cls["Accuracy"][1]:.3f}',  f'{cls["Accuracy"][1]-cls["Accuracy"][0]:+.3f}'],
        ['Precision', f'{cls["Precision"][0]:.3f}', f'{cls["Precision"][1]:.3f}', f'{cls["Precision"][1]-cls["Precision"][0]:+.3f}'],
        ['Recall',    f'{cls["Recall"][0]:.3f}',    f'{cls["Recall"][1]:.3f}',    f'{cls["Recall"][1]-cls["Recall"][0]:+.3f}'],
        ['F1-Score',  f'{cls["F1"][0]:.3f}',        f'{cls["F1"][1]:.3f}',        f'{cls["F1"][1]-cls["F1"][0]:+.3f}'],
        ['AUC-ROC',   f'{cls["AUC"][0]:.3f}',       f'{cls["AUC"][1]:.3f}',       f'{cls["AUC"][1]-cls["AUC"][0]:+.3f}'],
        ['Melhor algoritmo', cls['modelo_v1'][0],   cls['modelo_v2'][1],          '—'],
    ], headers=['Métrica', 'V1', 'V2', 'Δ'])

    add_para_just(doc, '')
    p = doc.add_paragraph()
    p.add_run('Leitura: ').bold = True
    p.add_run('a classificação se mantém praticamente idêntica em F1 e accuracy. '
              'O único movimento relevante é a queda do AUC-ROC (-0,02), o que '
              'reflete a perda de poder discriminativo nos extremos: na V2 o '
              'modelo já não tem mais os "casos óbvios" (NPS=0 com 5 dias de '
              'atraso e 7 reclamações) que eram facilmente capturados. Ainda '
              'assim, AUC=0,90 continua sendo um patamar excelente.')

    if (FIG_DIR / '06_comparativo_metricas_modelos.png').exists():
        doc.add_picture(str(FIG_DIR / '06_comparativo_metricas_modelos.png'),
                        width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ==================== 6. CONCLUSÃO E DECISÃO ====================
    doc.add_page_break()
    add_heading(doc, '6. Conclusão — Qual versão adotar?', 1)

    add_heading(doc, '6.1. O que aprendemos', 2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('A estrutura do problema é estável').bold = True
    p.add_run(': nenhuma correlação inverteu de sinal, os drivers principais '
              '(atraso, reclamação, recompra, CSAT) seguem sendo os mesmos. '
              'Isso significa que ')
    p.add_run('os NPS=0 não estão "distorcendo" a análise').italic = True
    p.add_run(' — apenas amplificam um padrão que já existe em toda a base.')

    p = doc.add_paragraph(style='List Bullet')
    p.add_run('A remoção mexe pouco no NPS oficial').bold = True
    p.add_run(' (-80,0 → -78,6). Mesmo "limpando" os 158 piores casos, a '
              'situação continua crítica. Confirma que o problema é sistêmico '
              'e não está concentrado em uma minoria de detratores extremos.')

    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Os modelos têm performance semelhante em ambas as versões').bold = True
    p.add_run('. A V2 não traz ganho substancial, e em algumas métricas tem '
              'leve perda (R² da regressão, AUC da classificação).')

    add_heading(doc, '6.2. Recomendação: manter a V1 como modelo de produção', 2)
    add_para_just(doc,
        'A V1 (com NPS=0) é a versão recomendada para uso em produção pelos '
        'seguintes motivos:'
    )
    p = doc.add_paragraph(style='List Number')
    p.add_run('Fidelidade ao negócio').bold = True
    p.add_run(': clientes que dão NPS=0 existem no mundo real e a empresa '
              'precisa identificá-los e tratá-los proativamente. Removê-los do '
              'treino faz o modelo perder sensibilidade exatamente para os '
              'casos mais críticos.')
    p = doc.add_paragraph(style='List Number')
    p.add_run('Sem ganho relevante de performance').bold = True
    p.add_run(': os modelos da V2 não são significativamente melhores e em '
              'algumas métricas são piores.')
    p = doc.add_paragraph(style='List Number')
    p.add_run('Risco de viés de seleção').bold = True
    p.add_run(': um modelo treinado sem detratores extremos pode subestimar '
              'sistematicamente o risco de NPS baixo em produção, gerando '
              'falsos negativos justamente nos casos onde a ação preventiva é '
              'mais valiosa.')

    add_heading(doc, '6.3. Quando a V2 seria útil?', 2)
    add_para_just(doc,
        'O exercício da V2 não é descartável — ele tem valor analítico em duas '
        'situações específicas:'
    )
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Validação de robustez').bold = True
    p.add_run(': mostrar à liderança que as conclusões não dependem dos '
              'extremos é um argumento forte contra críticas do tipo "esses '
              'NPS=0 estão distorcendo tudo".')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Análise de "clientes recuperáveis"').bold = True
    p.add_run(': se o objetivo for criar um modelo focado em mover clientes do '
              'meio (NPS 3-6) para Neutros/Promotores, faz sentido excluir os '
              'casos perdidos (NPS=0) e treinar especificamente nessa fatia '
              'da base.')

    add_para_just(doc, '')
    add_para_just(doc,
        'Para o objetivo central deste case — antecipar detratores e acionar '
        'fluxos preventivos — a V1 continua sendo a versão de referência.'
    )

    # Salvar
    output_path = REPORT_DIR / 'comparativo_v1_v2.docx'
    doc.save(output_path)
    print(f'Documento salvo: {output_path}')


if __name__ == '__main__':
    gerar()
